import openpyxl
from docx import Document
import re

def extract_and_manipulate_data(input_excel_path):
    wb = openpyxl.load_workbook(input_excel_path, data_only=True)
    sheet = None
    target_sheet_names = ['word_automation', 'Sheet1', 'sheet1']
    
    for sheet_name in target_sheet_names:
        if sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            break
    
    if sheet is None:
        sheet = wb.active
    
    processed_data = {}
    headers = [str(cell.value).strip() if cell.value is not None else '' for cell in sheet[1]]
    column_mapping = {
        'date': 'date',
        'Type of Audit': 'Type_of_Audit',
        'Type of Audit Report': 'Type_of_Audit_Report',
        'Period': 'Period',
        'Organization Name': 'Organization_Name',
        'Start Date': 'Start_Date',
        'End Date': 'End_Date',
        'Application Name/IP': 'Application_Name_IP',
        'From Date': 'From_Date',
        'Till Date': 'Till_Date'
    }
    
    column_indices = {}
    for excel_header, dict_key in column_mapping.items():
        try:
            index = next(
                i for i, header in enumerate(headers) 
                if excel_header.lower() in header.lower()
            )
            column_indices[dict_key] = index
        except StopIteration:
            pass
    
    for row_num, row in enumerate(sheet.iter_rows(min_row=2, values_only=True), start=2):
        row_data = {dict_key: str(row[col_index] or '').strip() for dict_key, col_index in column_indices.items() if col_index < len(row)}
        key = row_data.get('date', str(row_num))
        processed_data[key] = row_data
    
    return processed_data

def fill_word_template(template_path, output_path, processed_data, tables_excel_path):
    doc = Document(template_path)
    tables_wb = openpyxl.load_workbook(tables_excel_path, data_only=True)
    
    for paragraph in doc.paragraphs:
        if '{{' in paragraph.text and '}}' in paragraph.text:
            for key, value in processed_data.items():
                paragraph.text = re.sub(
                    r'\{\{data\.' + re.escape(key) + r'\.(\w+)\}\}', 
                    lambda m: str(value.get(m.group(1), '')), 
                    paragraph.text
                )
    
    for table in doc.tables:
        matched = False
        for sheet_name in tables_wb.sheetnames:
            sheet = tables_wb[sheet_name]
            sheet_headers = [str(cell.value).lower().strip() for cell in sheet[1]]
            table_headers = [cell.text.lower().strip() for cell in table.rows[0].cells if cell.text.strip()]
            
            if all(header in sheet_headers for header in table_headers):
                matched = True
                sheet_data = list(sheet.iter_rows(min_row=2, values_only=True))
                while len(table.rows) > 1:
                    table._element.remove(table.rows[-1]._element)
                
                for row_data in sheet_data:
                    row_cells = table.add_row().cells
                    for i, cell in enumerate(row_cells):
                        if i < len(sheet_headers):
                            cell.text = str(row_data[i]) if row_data[i] is not None else ''
                break
        
        if not matched:
            continue
    
    doc.save(output_path)

def main():
    input_excel_path = 'word_automation.xlsm'
    tables_excel_path = 'table_data.xlsx'
    template_path = 'generated_docs/document_1.docx'
    output_path = 'filled_document.docx'
    
    try:
        processed_data = extract_and_manipulate_data(input_excel_path)
        fill_word_template(template_path, output_path, processed_data, tables_excel_path)
        print("Document processing completed successfully!")
    except Exception as e:
        print(f"An error occurred: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()
