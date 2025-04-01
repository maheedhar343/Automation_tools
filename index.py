import os
import re
import subprocess
from docx import Document
import pandas as pd

def populate_word_template(template_path, data_dict, output_path):
    doc = Document(template_path)
    
    def replace_placeholders(text):
        if text is None:
            return None
        placeholders = re.findall(r'{{([^{}]+)}}', text)
        for placeholder in placeholders:
            if placeholder in data_dict:
                value = data_dict[placeholder]
                value = "" if pd.isna(value) else str(value)
                text = text.replace('{{' + placeholder + '}}', value)
            else:
                print(f"Warning: Placeholder '{placeholder}' not found in data")
        return text
    
    for paragraph in doc.paragraphs:
        if '{{' in paragraph.text and '}}' in paragraph.text:
            new_text = replace_placeholders(paragraph.text)
            if len(paragraph.runs) > 0:
                first_run = paragraph.runs[0]
                for run in paragraph.runs[1:]:
                    run.clear()
                first_run.text = new_text
            else:
                paragraph.add_run(new_text)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if '{{' in paragraph.text and '}}' in paragraph.text:
                        new_text = replace_placeholders(paragraph.text)
                        if len(paragraph.runs) > 0:
                            first_run = paragraph.runs[0]
                            for run in paragraph.runs[1:]:
                                run.clear()
                            first_run.text = new_text
                        else:
                            paragraph.add_run(new_text)
    
    doc.save(output_path)
    return output_path

def process_excel_to_word(template_path, excel_path, output_folder, filename_column=None):
    os.makedirs(output_folder, exist_ok=True)
    df = pd.read_excel(excel_path)
    generated_files = []
    
    for index, row in df.iterrows():
        data_dict = row.to_dict()
        base_filename = f"{data_dict[filename_column]}" if filename_column and filename_column in data_dict else f"document_{index+1}"
        base_filename = re.sub(r'[\\/*?:"<>|]', "", str(base_filename))
        output_path = os.path.join(output_folder, f"{base_filename}.docx")
        generated_file = populate_word_template(template_path, data_dict, output_path)
        generated_files.append(generated_file)
        print(f"Generated document: {output_path}")
    
    return generated_files

def run_table_update_script():
    try:
        print("Running table_update.py...")
        subprocess.run(["python", "table_update.py"], check=True)
        print("table_update.py executed successfully.")
    except subprocess.CalledProcessError as e:
        print(f"Error running table_update.py: {e}")

if __name__ == "__main__":
    template_path = "template.docx"
    excel_path = "data.xlsx"
    output_folder = "generated_docs"
    filename_column = "ID"
    
    generated_files = process_excel_to_word(template_path, excel_path, output_folder, filename_column)
    print(f"Total documents generated: {len(generated_files)}")
    print(f"Documents saved in folder: {output_folder}")
    
    if generated_files:
        run_table_update_script()
